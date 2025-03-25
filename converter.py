from pathlib import Path
from PIL import Image
import fitz
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import json
import csv
import xml.etree.ElementTree as ET


class Converter:
    def __init__(self, input_path: str, output_path: str, output_format: str) -> None:
        self.input_path = Path(input_path)
        self.output_path = Path(output_path)
        self.output_format = output_format.lower()

    def convert_image(self) -> None:
        format_map = {
            'png': "PNG",
            'jpg': "JPEG",
            'jpeg': "JPEG",
            'bmp': "BMP",
            'gif': "GIF"
        }
        if self.output_format not in format_map:
            raise ValueError(f"Unsupported image output format: {self.output_format}")
        with Image.open(self.input_path) as img:
            if self.output_format in ['jpg', 'jpeg'] and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(self.output_path, format=format_map[self.output_format])

    def convert_pdf(self) -> None:
        if self.output_format == 'txt':
            self.convert_pdf_to_txt()
        elif self.output_format == 'png':
            self.convert_pdf_to_png()
        elif self.output_format == 'docx':
            self.convert_pdf_to_docx()
        else:
            raise ValueError(f"Unsupported PDF output format: {self.output_format}")

    def convert_pdf_to_txt(self) -> None:
        doc = fitz.open(str(self.input_path))
        text = "\n".join(page.get_text() for page in doc)
        self.output_path.write_text(text, encoding='utf-8')

    def convert_pdf_to_png(self) -> None:
        doc = fitz.open(str(self.input_path))
        page = doc.load_page(0)
        pix = page.get_pixmap()
        pix.save(str(self.output_path))

    def convert_pdf_to_docx(self) -> None:
        doc = fitz.open(str(self.input_path))
        word_doc = Document()
        for page_num, page in enumerate(doc):
            text = page.get_text()
            word_doc.add_heading(f"Page {page_num + 1}", level=2)
            word_doc.add_paragraph(text)
            word_doc.add_paragraph("-" * 50)
        word_doc.save(str(self.output_path))

    def convert_docx(self) -> None:
        if self.output_format == 'txt':
            self.convert_docx_to_txt()
        elif self.output_format == 'pdf':
            self.convert_docx_to_pdf()
        else:
            raise ValueError(f"Unsupported DOCX output format: {self.output_format}")

    def convert_docx_to_txt(self) -> None:
        doc = Document(str(self.input_path))
        text = "\n".join(para.text for para in doc.paragraphs)
        self.output_path.write_text(text, encoding='utf-8')

    def convert_docx_to_pdf(self) -> None:
        doc = Document(str(self.input_path))
        c = canvas.Canvas(str(self.output_path), pagesize=letter)
        width, height = letter
        margin = 50
        y_position = height - margin
        line_height = 15
        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue
            # to check if there's enough space; if not, add a new page
            if y_position < margin + line_height:
                c.showPage()
                y_position = height - margin
            c.drawString(margin, y_position, text)
            y_position -= line_height
        c.save()

    def convert_csv(self) -> None:
        df = pd.read_csv(str(self.input_path))
        if self.output_format == 'json':
            df.to_json(str(self.output_path), orient='records', lines=True)
        elif self.output_format == 'txt':
            df.to_csv(str(self.output_path), index=False, sep='\t')
        else:
            raise ValueError(f"Unsupported CSV output format: {self.output_format}")

    def convert_text(self) -> None:
        if self.output_format == 'json':
            self.convert_text_to_json()
        elif self.output_format == 'csv':
            self.convert_text_to_csv()
        elif self.output_format == 'xml':
            self.convert_text_to_xml()
        elif self.output_format == 'docx':
            self.convert_text_to_docx()
        else:
            raise ValueError(f"Unsupported text output format: {self.output_format}")

    def convert_text_to_json(self) -> None:
        lines = self.input_path.read_text(encoding='utf-8').splitlines()
        data = [{"line": line} for line in lines]
        with self.output_path.open('w', encoding='utf-8') as json_file:
            json.dump(data, json_file, indent=4)

    def convert_text_to_csv(self) -> None:
        lines = self.input_path.read_text(encoding='utf-8').splitlines()
        rows = [line.split() for line in lines]
        with self.output_path.open('w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerows(rows)

    def convert_text_to_xml(self) -> None:
        root = ET.Element("root")
        lines = self.input_path.read_text(encoding='utf-8').splitlines()
        for idx, line in enumerate(lines, start=1):
            item = ET.SubElement(root, "item", id=str(idx))
            item.text = line
        tree = ET.ElementTree(root)
        tree.write(str(self.output_path), encoding='utf-8', xml_declaration=True)

    def convert_text_to_docx(self) -> None:
        doc = Document()
        lines = self.input_path.read_text(encoding='utf-8').splitlines()
        for line in lines:
            doc.add_paragraph(line)
        doc.save(str(self.output_path))

    def convert(self) -> None:
        input_format = self.input_path.suffix.lower().lstrip('.')
        if input_format in ['jpg', 'jpeg', 'png', 'bmp', 'gif']:
            self.convert_image()
        elif input_format == 'pdf':
            self.convert_pdf()
        elif input_format == 'docx':
            self.convert_docx()
        elif input_format == 'csv':
            self.convert_csv()
        elif input_format == 'txt':
            self.convert_text()
        else:
            raise ValueError(f"Unsupported input format: {input_format}")
